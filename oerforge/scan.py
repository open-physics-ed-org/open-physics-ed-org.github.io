"""
scan.py
--------
Static Site Asset and Content Scanner

This module provides functions for scanning site content, extracting assets, and populating the SQLite database
with page, section, and file records. It supports Markdown, Jupyter Notebooks, and DOCX files, and maintains
hierarchical relationships from the Table of Contents (TOC) YAML. All functions are documented and organized for
clarity and maintainability. Logging is standardized and all major operations are traced for debugging.

Key Features:
- Batch reading and asset extraction for supported file types
- Hierarchical TOC walking and database population
- Asset linking and MIME type detection
- Section and descendant queries using recursive CTEs

Usage:
    Import and call scan_toc_and_populate_db(config_path) to scan the TOC and populate the database.
    Use get_descendants_for_parent() to query section hierarchies.
"""

import os
import sqlite3
import re
import logging

import logging
import os

project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
build_log_path = os.path.join(project_root, 'log', 'build.log')
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s %(levelname)s %(message)s',
    handlers=[
        logging.FileHandler(build_log_path, mode='a', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# =========================
# File Reading Utilities
# =========================

def batch_read_files(file_paths):
    """
    Reads multiple files and returns their contents as a dict: {path: content}
    Supports markdown (.md), notebook (.ipynb), docx (.docx), and other file types.
    """
    contents = {}
    for path in file_paths:
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == '.md':
                contents[path] = read_markdown_file(path)
            elif ext == '.ipynb':
                contents[path] = read_notebook_file(path)
            elif ext == '.docx':
                contents[path] = read_docx_file(path)
            else:
                contents[path] = None
        except Exception as e:
            logging.error(f"Could not read {path}: {e}")
            contents[path] = None
    return contents

def read_markdown_file(path):
    """
    Reads a markdown (.md) file and returns its content as a string.
    """
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logging.error(f"Could not read markdown file {path}: {e}")
        return None

def read_notebook_file(path):
    """
    Reads a Jupyter notebook (.ipynb) file and returns its content as a dict.
    """
    import json
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logging.error(f"Could not read notebook file {path}: {e}")
        return None

def read_docx_file(path):
    """
    Reads a docx file and returns its text content as a string.
    Requires python-docx to be installed.
    """
    try:
        from docx import Document
        doc = Document(path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)
    except ImportError:
        logging.error("python-docx is not installed. Run 'pip install python-docx' in your environment.")
        return None
    except Exception as e:
        logging.error(f"Could not read docx file {path}: {e}")
        return None

# ==========================
# Asset Extraction Utilities
# ==========================
def batch_extract_assets(contents_dict, content_type, **kwargs):
    """
    Extracts assets from multiple file contents in one pass.
    contents_dict: {path: content}
    content_type: 'markdown', 'notebook', 'docx', etc.
    Returns a dict: {path: [asset_records]}
    """
    from oerforge.db_utils import insert_records, link_files_to_pages, get_db_connection
    assets = {}
    # Helper: MIME type mapping (media, document, and data types)
    mime_map = {
        # Images
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.svg': 'image/svg+xml',
        '.bmp': 'image/bmp',
        '.webp': 'image/webp',
        # Video
        '.mp4': 'video/mp4',
        '.webm': 'video/webm',
        '.mov': 'video/quicktime',
        '.avi': 'video/x-msvideo',
        '.mkv': 'video/x-matroska',
        # Audio
        '.mp3': 'audio/mpeg',
        '.wav': 'audio/wav',
        '.ogg': 'audio/ogg',
        '.flac': 'audio/flac',
        # Documents
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        # Data files
        '.csv': 'text/csv',
        '.tsv': 'text/tab-separated-values',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.npy': 'application/octet-stream',
        '.txt': 'text/plain',
        '.zip': 'application/zip',
        '.tar': 'application/x-tar',
        '.gz': 'application/gzip',
        '.rst': 'text/x-rst',
        # Markdown/Notebook
        '.md': 'text/markdown',
        '.ipynb': 'application/x-ipynb+json',
    }
    # Insert each source file as a page if not present
    import threading
    import time
    conn = get_db_connection()
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Opened DB connection in batch_extract_assets at {time.time()}")
    cursor = conn.cursor()
    # Add mime_type column to content if not present
    cursor.execute("PRAGMA table_info(content)")
    columns = [row[1] for row in cursor.fetchall()]
    if 'mime_type' not in columns:
        try:
            cursor.execute("ALTER TABLE content ADD COLUMN mime_type TEXT")
        except Exception:
            pass
    for source_path in contents_dict:
        ext = os.path.splitext(source_path)[1].lower()
        mime_type = mime_map.get(ext, '')
        cursor.execute("SELECT id FROM content WHERE source_path=?", (source_path,))
        if cursor.fetchone() is None:
            cursor.execute("INSERT INTO content (source_path, output_path, is_autobuilt, mime_type) VALUES (?, ?, ?, ?)", (source_path, None, 0, mime_type))
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Committing DB in batch_extract_assets at {time.time()}")
    try:
        conn.commit()
    except Exception as e:
        import traceback
        logging.error(f"[ERROR][{os.getpid()}][{threading.get_ident()}] Commit failed in batch_extract_assets: {e}\n{traceback.format_exc()}")
        raise
    # Extract assets for each file type
    for path, content in contents_dict.items():
        ext = os.path.splitext(path)[1].lower()
        if content_type == 'markdown':
            assets[path] = [a for a in extract_linked_files_from_markdown_content(content, page_id=None)
                            if mime_map.get(os.path.splitext(a.get('path',''))[1].lower())]
        elif content_type == 'notebook':
            if content and isinstance(content, dict) and 'cells' in content:
                cell_assets = []
                for cell in content['cells']:
                    cell_assets.extend([a for a in extract_linked_files_from_notebook_cell_content(cell, nb_path=path)
                                       if mime_map.get(os.path.splitext(a.get('path',''))[1].lower())])
                assets[path] = cell_assets
            else:
                assets[path] = []
        elif content_type == 'docx':
            assets[path] = [a for a in extract_linked_files_from_docx_content(path, page_id=None)
                            if mime_map.get(os.path.splitext(a.get('path',''))[1].lower())] if content else []
        else:
            assets[path] = []
    # Insert asset records into files table
    file_records = []
    file_page_links = []
    for source_path, asset_list in assets.items():
        for asset in asset_list:
            asset_path = asset.get('path', '')
            asset_ext = os.path.splitext(asset_path)[1].lower()
            mime_type = mime_map.get(asset_ext, '')
            file_record = {
                'filename': os.path.basename(asset_path),
                'extension': asset_ext,
                'mime_type': mime_type,
                'is_image': int(asset_ext in ['.png','.jpg','.jpeg','.gif','.svg']),
                'is_remote': int(asset_path.startswith('http')),
                'url': asset_path,
                'referenced_page': source_path,
                'relative_path': asset_path,
                'absolute_path': None,
                'cell_type': asset.get('type', None),
                'is_code_generated': None,
                'is_embedded': None
            }
            file_records.append(file_record)
    file_ids = insert_records('files', file_records, conn=conn, cursor=cursor)
    # Link files to pages
    idx = 0
    for source_path, asset_list in assets.items():
        for _ in asset_list:
            file_page_links.append((file_ids[idx], source_path))
            idx += 1
    if file_page_links:
        link_files_to_pages(file_page_links, conn=conn, cursor=cursor)
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Closing DB connection in batch_extract_assets at {time.time()}")
    conn.close()
    return assets

def extract_linked_files_from_markdown_content(md_text, page_id=None):
    """
    Extract asset links from markdown text.
    Args:
        md_text (str): Markdown content.
        page_id (optional): Page identifier for DB linking.
    Returns:
        list: File record dicts for each asset found.
    """
    import re
    asset_pattern = re.compile(r'!\[[^\]]*\]\(([^)]+)\)|\[[^\]]*\]\(([^)]+)\)')
    assets = []
    for match in asset_pattern.finditer(md_text):
        asset_path = match.group(1) or match.group(2)
        if asset_path:
            assets.append({
                'type': 'asset',
                'path': asset_path,
                'page_id': page_id
            })
    return assets

def extract_linked_files_from_notebook_cell_content(cell, nb_path=None):
    """
    Extract asset links from a notebook cell.
    Args:
        cell (dict): Notebook cell.
        nb_path (str, optional): Notebook file path.
    Returns:
        list: File record dicts for each asset found.
    """
    assets = []
    # Extract markdown-linked images
    if cell.get('cell_type') == 'markdown':
        source = cell.get('source', [])
        if isinstance(source, list):
            text = ''.join(source)
        else:
            text = str(source)
        assets.extend(extract_linked_files_from_markdown_content(text, page_id=None))
        for asset in assets:
            asset['notebook'] = nb_path
    # Extract embedded/code-produced images from outputs
    if cell.get('cell_type') == 'code' and 'outputs' in cell:
        for idx, output in enumerate(cell['outputs']):
            # Typical image output: {'data': {'image/png': ...}, ...}
            if 'data' in output:
                for img_type in ['image/png', 'image/jpeg', 'image/gif', 'image/svg+xml']:
                    if img_type in output['data']:
                        ext = {
                            'image/png': '.png',
                            'image/jpeg': '.jpg',
                            'image/gif': '.gif',
                            'image/svg+xml': '.svg',
                        }[img_type]
                        nb_name = os.path.basename(nb_path) if nb_path else 'notebook'
                        rel_path = f'notebook_embedded/{nb_name}/cell{idx}{ext}'
                        assets.append({
                            'type': 'asset',
                            'path': rel_path,
                            'notebook': nb_path,
                            'filename': f'cell{idx}{ext}',
                            'extension': ext,
                            'is_embedded': True,
                            'is_code_generated': True
                        })
    return assets

def extract_linked_files_from_docx_content(docx_path, page_id=None):
    """
    Extract asset links from a DOCX file.
    Args:
        docx_path (str): Path to DOCX file.
        page_id (optional): Page identifier for DB linking.
    Returns:
        list: File record dicts for each asset found.
    """
    assets = []
    try:
        from docx import Document
        doc = Document(docx_path)
        import re
        asset_pattern = re.compile(r'(https?://[^\s]+|assets/[^\s]+|images/[^\s]+)')
        # Extract text-based links as before
        for para in doc.paragraphs:
            matches = asset_pattern.findall(para.text)
            for asset_path in matches:
                assets.append({
                    'type': 'asset',
                    'path': asset_path,
                    'page_id': page_id
                })
        # Extract embedded images
        for rel in doc.part.rels.values():
            if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
                img_part = rel.target_part
                img_name = os.path.basename(img_part.partname)
                img_ext = os.path.splitext(img_name)[1].lower()
                # Use a relative path for DB, e.g. 'docx_embedded/<docx_filename>/<img_name>'
                rel_path = f'docx_embedded/{os.path.basename(docx_path)}/{img_name}'
                assets.append({
                    'type': 'asset',
                    'path': rel_path,
                    'page_id': page_id,
                    'filename': img_name,
                    'extension': img_ext,
                    'is_embedded': True
                })
    except Exception as e:
        logging.error(f"Could not extract assets from docx {docx_path}: {e}")
    return assets

def populate_site_info_from_config(config_filename='_config.yml'):
    """
    Populate the site_info table from the given config file (default: _config.yml).
    Args:
        config_filename (str): Name of the config file (e.g., '_config.yml').
    """
    import yaml
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    db_path = os.path.join(project_root, 'db', 'sqlite.db')
    full_config_path = os.path.join(project_root, config_filename)
    with open(full_config_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)
    site = config.get('site', {})
    footer = config.get('footer', {})
    theme = site.get('theme', {})
    # Read header.html contents
    header_html_path = os.path.join(project_root, 'layouts', 'partials', 'header.html')
    try:
        with open(header_html_path, 'r', encoding='utf-8') as hf:
            header_html = hf.read()
    except Exception:
        header_html = ''
    logging.debug(f"[DEBUG] header_html read from file (first 500 chars): {repr(header_html)[:500]}")
    import threading
    import time
    conn = sqlite3.connect(db_path)
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Opened DB connection in populate_site_info_from_config at {time.time()}")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM site_info")
    cursor.execute(
        """
        INSERT INTO site_info (title, author, description, logo, favicon, theme_default, theme_light, theme_dark, language, github_url, footer_text, header)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            site.get('title', ''),
            site.get('author', ''),
            site.get('description', ''),
            site.get('logo', ''),
            site.get('favicon', ''),
            theme.get('default', ''),
            theme.get('light', ''),
            theme.get('dark', ''),
            site.get('language', ''),
            site.get('github_url', ''),
            footer.get('text', ''),
            header_html
        )
    )
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Committing DB in populate_site_info_from_config at {time.time()}")
    try:
        conn.commit()
    except Exception as e:
        import traceback
        logging.error(f"[ERROR][{os.getpid()}][{threading.get_ident()}] Commit failed in populate_site_info_from_config: {e}\n{traceback.format_exc()}")
        raise
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Closing DB connection in populate_site_info_from_config at {time.time()}")
    conn.close()

# ----
# Conversion Capability Helper
# ----

def get_conversion_flags(extension):
    # Local import to avoid circular import issues
    from oerforge.db_utils import get_enabled_conversions
    """
    Get conversion flags for a given file extension using the DB.
    Args:
        extension (str): File extension (e.g., '.md', '.ipynb').
    Returns:
        dict: Conversion capability flags.
    """
    targets = get_enabled_conversions(extension)
    # Map target formats to flags
    flag_map = {
        '.md': 'can_convert_md',
        '.tex': 'can_convert_tex',
        '.pdf': 'can_convert_pdf',
        '.docx': 'can_convert_docx',
        '.ppt': 'can_convert_ppt',
        '.jupyter': 'can_convert_jupyter',
        '.ipynb': 'can_convert_ipynb'
    }
    flags = {v: False for v in flag_map.values()}
    for t in targets:
        if t in flag_map:
            flags[flag_map[t]] = True
    return flags

def scan_toc_and_populate_db(config_path):
    """
    Walk the TOC from the config YAML, read each file, extract assets/images, and populate the DB with content and asset records.
    Maintains TOC hierarchy and section relationships.
    Args:
        config_path (str): Path to the config YAML file.
    """
    import yaml
    from oerforge.db_utils import get_db_connection, insert_records, link_files_to_pages
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    full_config_path = os.path.join(project_root, config_path)
    with open(full_config_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)
    toc = config.get('toc', [])

    import threading
    import time
    conn = get_db_connection()
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Opened DB connection in scan_toc_and_populate_db at {time.time()}")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM content")
    seen_paths = set()
    file_paths = []
    # Removed outdated import of insert_file_records; link_files_to_pages is already imported above

    from oerforge.db_utils import set_relative_link, set_menu_context
    def walk_toc(items, parent_output_path=None, parent_slug=None, parent_menu_context=None, level=0):
        content_records = []
        for idx, item in enumerate(items):
            file_path = item.get('file')
            title = item.get('title', None)
            order = int(idx)
            # Debug: log order and level for each item
            logging.debug(f"[DEBUG][walk_toc] idx={idx} title={title} file_path={file_path} item={item} order={order} (type={type(order)}) level={level} (type={type(level)})")
            item_slug = item.get('slug', re.sub(r'[^a-zA-Z0-9]+', '_', title.lower()).strip('_')) if title else f'section_{idx}'
            effective_slug = parent_slug if parent_slug else item_slug
            menu_context = item.get('menu_context', parent_menu_context if parent_menu_context else 'main')
            if file_path:
                source_path = file_path if file_path.startswith('content/') else f'content/{file_path}'
                ext = os.path.splitext(source_path)[1].lower()
                rel_path = source_path[8:] if source_path.startswith('content/') else source_path
                base_name = os.path.splitext(os.path.basename(rel_path))[0]
                parent_dir = os.path.basename(os.path.dirname(source_path))
                # Compute relative_link for menu (strip 'build/' and use output_path)
                if source_path == f'content/{base_name}.md' and base_name != 'index':
                    output_path = os.path.join('build', base_name, 'index.html')
                    output_path_debug = 'TOP-LEVEL-EXPLICIT'
                elif base_name == parent_dir:
                    output_path = os.path.join('build', effective_slug, 'index.html')
                    output_path_debug = 'SUBFOLDER-MATCH'
                else:
                    output_path = os.path.join('build', effective_slug, base_name + '.html')
                    output_path_debug = 'DEFAULT'
                # Compute relative_link for menu (strip 'build/' if present)
                relative_link = output_path[6:] if output_path.startswith('build/') else output_path
                logging.debug(f"[DEBUG][walk_toc] output_path chosen for '{title}': {output_path} (mode: {output_path_debug})")
                flags = get_conversion_flags(ext)
                content_record = {
                    'title': title,
                    'source_path': source_path,
                    'output_path': output_path,
                    'is_autobuilt': 0,
                    'mime_type': ext,
                    'can_convert_md': flags['can_convert_md'],
                    'can_convert_tex': flags['can_convert_tex'],
                    'can_convert_pdf': flags['can_convert_pdf'],
                    'can_convert_docx': flags['can_convert_docx'],
                    'can_convert_ppt': flags['can_convert_ppt'],
                    'can_convert_jupyter': flags['can_convert_jupyter'],
                    'can_convert_ipynb': flags['can_convert_ipynb'],
                    'parent_output_path': parent_output_path,
                    'slug': effective_slug,
                    'order': int(order),
                    'relative_link': relative_link,
                    'menu_context': menu_context,
                    'level': int(level)
                }
                logging.debug(f"[DEBUG][walk_toc] content_record={content_record}")
                content_records.append(content_record)
                abs_path = os.path.join(project_root, source_path)
                file_paths.append(abs_path)
                # Optionally set in DB after insert (not needed here, but available)
                children = item.get('children', [])
                if children:
                    child_records = walk_toc(children, parent_output_path=output_path, parent_slug=effective_slug, parent_menu_context=menu_context, level=int(level)+1)
                    content_records.extend(child_records)
            elif item.get('children'):
                output_path = os.path.join('build', item_slug, 'index.html')
                relative_link = output_path[6:] if output_path.startswith('build/') else output_path
                content_record = {
                    'title': title,
                    'source_path': None,
                    'output_path': output_path,
                    'is_autobuilt': 1,
                    'mime_type': 'section',
                    'can_convert_md': False,
                    'can_convert_tex': False,
                    'can_convert_pdf': False,
                    'can_convert_docx': False,
                    'can_convert_ppt': False,
                    'can_convert_jupyter': False,
                    'can_convert_ipynb': False,
                    'parent_output_path': parent_output_path,
                    'slug': item_slug,
                    'order': int(order),
                    'relative_link': relative_link,
                    'menu_context': 'main',  # Always set menu_context for section index
                    'level': int(level)
                }
                logging.debug(f"[DEBUG][walk_toc] content_record (section)={content_record}")
                content_records.append(content_record)
                children = item.get('children', [])
                if children:
                    child_records = walk_toc(children, parent_output_path=output_path, parent_slug=item_slug, parent_menu_context=menu_context, level=int(level)+1)
                    content_records.extend(child_records)
        return content_records

    # Usage in scan_toc_and_populate_db:
    all_content_records = walk_toc(toc)

    # Deduplicate by (source_path, output_path, title)
    unique_records = {}
    for rec in all_content_records:
        key = (rec.get('source_path'), rec.get('output_path'), rec.get('title'))
        if key not in unique_records:
            unique_records[key] = rec
    deduped_records = list(unique_records.values())

    # Debug: log all deduped records' order and level, and show their types
    for rec in deduped_records:
        logging.debug(f"[DEBUG][deduped] title={rec.get('title')} order={rec.get('order')} (type={type(rec.get('order'))}) level={rec.get('level')} (type={type(rec.get('level'))})")

    insert_records('content', deduped_records, db_path=os.path.join(project_root, 'db', 'sqlite.db'), conn=conn, cursor=cursor)
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Committing DB in scan_toc_and_populate_db at {time.time()}")
    try:
        conn.commit()
    except Exception as e:
        import traceback
        logging.error(f"[ERROR][{os.getpid()}][{threading.get_ident()}] Commit failed in scan_toc_and_populate_db: {e}\n{traceback.format_exc()}")
        raise

    # Read all files and extract assets
    rel_file_paths = [os.path.relpath(p, project_root) for p in file_paths if os.path.exists(p)]
    contents = batch_read_files(rel_file_paths)
    for path in rel_file_paths:
        ext = os.path.splitext(path)[1].lower()
        if ext == '.md':
            batch_extract_assets({path: contents[path]}, 'markdown', conn=conn, cursor=cursor)
        elif ext == '.ipynb':
            batch_extract_assets({path: contents[path]}, 'notebook', conn=conn, cursor=cursor)
        elif ext == '.docx':
            batch_extract_assets({path: contents[path]}, 'docx', conn=conn, cursor=cursor)
        # Add more types as needed
    logging.debug(f"[DEBUG][{os.getpid()}][{threading.get_ident()}] Closing DB connection in scan_toc_and_populate_db at {time.time()}")
    conn.close()

# ----
# Recursive CTE Helper for Section Index Generation
# ----
def get_descendants_for_parent(parent_output_path, db_path):
    """
    Query all children, grandchildren, and deeper descendants for a given parent_output_path using a recursive CTE.
    Args:
        parent_output_path (str): Output path of the parent section.
        db_path (str): Path to the SQLite database.
    Returns:
        list: Dicts for each descendant (id, title, output_path, parent_output_path, slug, level).
    """
    import sqlite3
    import logging
    logging.info(f"[DB-OPEN] Attempting to open database: {db_path}")
    conn = sqlite3.connect(db_path)
    logging.info(f"[DB-OPEN] Database connection established: {db_path}")
    cursor = conn.cursor()
    query = '''
    WITH RECURSIVE content_hierarchy(id, title, output_path, parent_output_path, slug, level) AS (
      SELECT id, title, output_path, parent_output_path, slug, 0 as level
      FROM content
      WHERE output_path = ?
      UNION ALL
      SELECT c.id, c.title, c.output_path, c.parent_output_path, c.slug, ch.level + 1
      FROM content c
      JOIN content_hierarchy ch ON c.parent_output_path = ch.output_path
    )
    SELECT id, title, output_path, parent_output_path, slug, level FROM content_hierarchy WHERE level > 0 ORDER BY level, output_path;
    '''
    cursor.execute(query, (parent_output_path,))
    rows = cursor.fetchall()
    logging.info(f"[DB-CLOSE] Database connection closed: {db_path}")
    conn.close()
    # Return as list of dicts
    return [
        {
            'id': row[0],
            'title': row[1],
            'output_path': row[2],
            'parent_output_path': row[3],
            'slug': row[4],
            'level': row[5]
        }
        for row in rows
    ]

if __name__ == "__main__":
    # Default config file name
    config_file = "_content.yml"
    import sys
    if len(sys.argv) > 1:
        config_file = sys.argv[1]
    logging.info(f"[MAIN] Running scan_toc_and_populate_db with config: {config_file}")
    scan_toc_and_populate_db(config_file)